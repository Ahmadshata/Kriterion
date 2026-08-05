"""Text extraction utilities for PDF and DOCX resume files."""

import re
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import fitz
from docx import Document as DocxDocument


# -----------------------------
# Text extraction
# -----------------------------


def extract_text_by_page(pdf_path: Path) -> Tuple[List[str], bool]:
    doc = fitz.open(pdf_path)
    pages_standard = [page.get_text("text") for page in doc]
    used_ocr = False

    text = "\n".join(pages_standard).strip()

    # If text is too short, try OCR
    if len(text) < 500:
        ocr = try_ocr(doc)
        if ocr:
            return ocr, True

    # Detect multi-column and use block-based extraction if needed
    if _is_multi_column(doc):
        pages_blocks = _extract_with_blocks(doc)
        return pages_blocks, False

    return pages_standard, used_ocr


def extract_text_from_docx(docx_path: Path) -> Tuple[List[str], bool]:
    """Extract text from a DOCX file. Returns (pages_text, used_ocr=False)."""
    doc = DocxDocument(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # DOCX doesn't have page breaks reliably — treat entire doc as one page
    full_text = "\n".join(paragraphs)

    # Also extract text from tables (common in CVs)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                full_text += "\n" + row_text

    return [full_text], False


def _extract_with_blocks(doc: fitz.Document) -> List[str]:
    """Use get_text('blocks') sorted by (y_band, x) for multi-column layouts."""
    pages: List[str] = []
    for page in doc:
        blocks = page.get_text("blocks")
        text_blocks = [b for b in blocks if b[6] == 0]
        # Sort by y-position (10pt tolerance band), then x-position
        text_blocks.sort(key=lambda b: (round(b[1] / 10) * 10, b[0]))
        page_text = "\n".join(b[4].strip() for b in text_blocks if b[4].strip())
        pages.append(page_text)
    return pages


def _is_multi_column(doc: fitz.Document) -> bool:
    """Heuristic: detect if first content page has multi-column layout."""
    if len(doc) == 0:
        return False
    page = doc[0]
    blocks = page.get_text("blocks")
    text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]
    if len(text_blocks) < 4:
        return False

    y_bands: Dict[int, List[float]] = {}
    for b in text_blocks:
        band = round(b[1] / 10)
        y_bands.setdefault(band, []).append(b[0])

    page_width = page.rect.width
    multi_col_bands = 0
    for x_positions in y_bands.values():
        if len(x_positions) >= 2:
            x_spread = max(x_positions) - min(x_positions)
            if x_spread > page_width * 0.3:
                multi_col_bands += 1

    return multi_col_bands >= 3


def try_ocr(doc: fitz.Document) -> Optional[List[str]]:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception:
        return None

    ocr_pages: List[str] = []
    for page in doc:
        pix = page.get_pixmap(dpi=200)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        ocr_pages.append(pytesseract.image_to_string(img))
    return ocr_pages


_BULLET_RE = re.compile(r"^[•‣◦⁃∙●○▪▫–—•·‣⁃►▪▸-]\s*$")


def _merge_bullet_lines(lines: List[str]) -> List[str]:
    """Merge standalone bullet-character lines with the line that follows."""
    merged: List[str] = []
    i = 0
    while i < len(lines):
        if _BULLET_RE.match(lines[i]) and i + 1 < len(lines):
            merged.append(lines[i].rstrip() + " " + lines[i + 1].lstrip())
            i += 2
        else:
            merged.append(lines[i])
            i += 1
    return merged


def iter_lines_with_pages(pages: Sequence[str]) -> Iterable[Tuple[int, str]]:
    for i, page_text in enumerate(pages, start=1):
        raw_lines = page_text.splitlines()
        for line in _merge_bullet_lines(raw_lines):
            yield i, line.rstrip()
